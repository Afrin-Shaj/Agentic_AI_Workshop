import streamlit as st
import os
import logging
from PyPDF2 import PdfReader
import re
import json
from typing import Dict, List
import pandas as pd
import numpy as np
from dotenv import load_dotenv
import requests
from bs4 import BeautifulSoup
import time
from io import BytesIO
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
from sentence_transformers import SentenceTransformer
import faiss
import google.generativeai as genai
from langchain.text_splitter import RecursiveCharacterTextSplitter
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from urllib.parse import urlparse
import jsonschema

# Configure logging
logging.basicConfig(level=logging.INFO, filename='culture_fit_analyzer.log',
                    format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Load environment variables
load_dotenv()
api_key = os.getenv("GOOGLE_API_KEY")
linkedin_email = os.getenv("LINKEDIN_EMAIL")
linkedin_password = os.getenv("LINKEDIN_PASSWORD")

if not api_key:
    st.error("GOOGLE_API_KEY not found in .env file.")
    st.stop()
if not all([linkedin_email, linkedin_password]):
    st.error("LinkedIn credentials not found in .env file.")
    st.stop()

# Initialize LLM and Embedding Model
genai.configure(api_key=api_key)
llm = genai.GenerativeModel("gemini-1.5-flash", generation_config={"temperature": 0.3, "max_output_tokens": 4096})
embedder = SentenceTransformer('all-MiniLM-L6-v2')
text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)

# Initialize FAISS vector store and embedding cache
dimension = 384
index = faiss.IndexFlatL2(dimension)
chunk_id_to_text = {}
embedding_cache = {}  # Cache for query embeddings

# JSON schemas for agent outputs
AGENT_SCHEMAS = {
    "candidate_behavior": {
        "type": "object",
        "properties": {
            "personality_traits": {"type": "array", "items": {"type": "object"}},
            "soft_skills": {"type": "array", "items": {"type": "object"}},
            "communication_patterns": {"type": "array", "items": {"type": "object"}},
            "work_style": {"type": "array", "items": {"type": "object"}},
            "core_values": {"type": "array", "items": {"type": "object"}},
            "emotional_intelligence": {"type": "string"},
            "summary": {"type": "string"}
        },
        "required": ["personality_traits", "soft_skills", "summary"]
    },
    "company_culture": {
        "type": "object",
        "properties": {
            "core_values": {"type": "array", "items": {"type": "object"}},
            "work_environment": {"type": "array", "items": {"type": "object"}},
            "behavioral_expectations": {"type": "array", "items": {"type": "object"}},
            "communication_norms": {"type": "array", "items": {"type": "object"}},
            "preferred_work_mode": {"type": "string"},
            "cultural_red_flags": {"type": "array", "items": {"type": "object"}},
            "summary": {"type": "string"}
        },
        "required": ["core_values", "work_environment", "summary"]
    },
    "compatibility_mapping": {
        "type": "object",
        "properties": {
            "compatibility_scores": {"type": "object"},
            "alignment_strengths": {"type": "array", "items": {"type": "object"}},
            "potential_mismatches": {"type": "array", "items": {"type": "object"}},
            "adaptation_requirements": {"type": "array", "items": {"type": "object"}},
            "summary": {"type": "string"}
        },
        "required": ["compatibility_scores", "summary"]
    },
    "culture_fit_scoring": {
        "type": "object",
        "properties": {
            "overall_fit_score": {"type": "number"},
            "detailed_scores": {"type": "object"},
            "strengths": {"type": "array", "items": {"type": "object"}},
            "challenges": {"type": "array", "items": {"type": "object"}},
            "recommendations": {"type": "array", "items": {"type": "object"}},
            "interview_questions": {"type": "array", "items": {"type": "object"}},
            "onboarding_plan": {"type": "array", "items": {"type": "object"}},
            "summary": {"type": "string"}
        },
        "required": ["overall_fit_score", "detailed_scores", "summary"]
    }
}

def scrape_website_content(url: str) -> str:
    """Scrape content from a website URL, using Selenium for LinkedIn"""
    try:
        parsed_url = urlparse(url)
        clean_url = f"{parsed_url.scheme}://{parsed_url.netloc}{parsed_url.path}"
        
        if 'linkedin.com' in parsed_url.netloc:
            chrome_options = Options()
            chrome_options.add_argument("--no-sandbox")
            chrome_options.add_argument("--disable-dev-shm-usage")
            chrome_options.add_argument("user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/117.0.0.0 Safari/537.36")
            
            driver = webdriver.Chrome(options=chrome_options)
            try:
                driver.get("https://www.linkedin.com/login")
                WebDriverWait(driver, 10).until(
                    EC.presence_of_element_located((By.ID, "username"))
                )
                
                driver.find_element(By.ID, "username").send_keys(linkedin_email)
                driver.find_element(By.ID, "password").send_keys(linkedin_password)
                driver.find_element(By.XPATH, "//button[@type='submit']").click()
                
                try:
                    WebDriverWait(driver, 15).until(
                        EC.url_contains("feed")
                    )
                except:
                    if driver.find_elements(By.ID, "input__phone_verification_pin") or "verify" in driver.current_url.lower():
                        logger.warning("MFA required for LinkedIn login")
                        st.error("LinkedIn login requires MFA. Use manual input in the 'Manual Input' tab.")
                        return ""
                    else:
                        logger.warning("LinkedIn login failed")
                        st.error("LinkedIn login failed. Verify credentials or use manual input.")
                        return ""
                
                driver.get(clean_url)
                WebDriverWait(driver, 15).until(
                    EC.presence_of_element_located((By.TAG_NAME, "main"))
                )
                
                if "signin" in driver.current_url or driver.find_elements(By.CLASS_NAME, "authwall"):
                    logger.warning(f"Access restricted for LinkedIn URL: {clean_url}")
                    st.error("Unable to access LinkedIn profile (private/restricted). Use manual input.")
                    return ""
                
                soup = BeautifulSoup(driver.page_source, 'html.parser')
                profile_content = soup.find('main')
                if not profile_content:
                    logger.warning(f"No main content found for {clean_url}")
                    st.error("Failed to extract profile content. Try manual input.")
                    return ""
                
                text = profile_content.get_text(separator=' ', strip=True)
                text = re.sub(r'\s+', ' ', text).strip()[:6000]
                logger.info(f"Successfully scraped LinkedIn profile: {clean_url}")
                return text
            
            except Exception as e:
                logger.error(f"Selenium error scraping LinkedIn {clean_url}: {str(e)}")
                st.error(f"Error scraping LinkedIn profile: {str(e)}. Try manual input.")
                return ""
            finally:
                driver.quit()
        
        else:
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/117.0.0.0 Safari/537.36',
                'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8'
            }
            response = requests.get(clean_url, headers=headers, timeout=10, allow_redirects=True)
            response.raise_for_status()
            soup = BeautifulSoup(response.content, 'html.parser')
            for script in soup(["script", "style", "nav", "footer"]):
                script.extract()
            text = soup.get_text(separator=' ', strip=True)
            text = re.sub(r'\s+', ' ', text).strip()
            if any(keyword in text.lower() for keyword in ['sign in', 'join now', 'password']):
                logger.warning(f"Login page content detected for {clean_url}")
                st.error(f"Scraped content is a login page. Use manual input.")
                return ""
            if len(text) > 6000:
                text = text[:6000] + "..."
            logger.info(f"Successfully scraped content from {clean_url}")
            return text
            
    except Exception as e:
        logger.error(f"Error scraping {clean_url}: {str(e)}")
        st.error(f"Error scraping {clean_url}: {str(e)}")
        return ""

def extract_text_from_pdf(uploaded_file) -> str:
    """Extract text from PDF"""
    try:
        reader = PdfReader(uploaded_file)
        text = "".join(page.extract_text() or "" for page in reader.pages)
        logger.info("Successfully extracted text from PDF")
        return clean_text(text)
    except Exception as e:
        logger.error(f"Error reading PDF: {str(e)}")
        st.error(f"Error reading PDF: {str(e)}")
        return ""

def clean_text(text: str) -> str:
    """Clean and normalize text"""
    if not text:
        return ""
    text = re.sub(r'[^\w\s\-.,;:!?()]', '', text)
    text = re.sub(r'\s+', ' ', text.strip())
    return text

def create_chunks_and_embeddings(text: str, source: str) -> List[Dict]:
    """Split text into chunks and create embeddings for RAG"""
    if not text:
        return []
    chunks = text_splitter.split_text(text)
    chunk_data = []
    for i, chunk in enumerate(chunks):
        chunk_id = f"{source}_{i}"
        chunk_data.append({"id": chunk_id, "text": chunk, "source": source})
        chunk_id_to_text[chunk_id] = chunk
    embeddings = embedder.encode([chunk['text'] for chunk in chunk_data], convert_to_numpy=True)
    index.add(embeddings)
    logger.info(f"Created {len(chunks)} chunks and embeddings for {source}")
    return chunk_data

def retrieve_relevant_chunks(query: str, k: int = 5) -> List[str]:
    """Retrieve top-k relevant chunks using FAISS for RAG with dynamic k"""
    try:
        # Use cached embedding if available
        if query in embedding_cache:
            query_embedding = embedding_cache[query]
        else:
            query_embedding = embedder.encode([query], convert_to_numpy=True)
            embedding_cache[query] = query_embedding
        
        # Dynamically adjust k based on data volume
        k_dynamic = min(k, max(1, len(chunk_id_to_text) // 2 or 10))
        distances, indices = index.search(query_embedding, k_dynamic)
        retrieved_chunks = []
        for idx in indices[0]:
            if idx < len(chunk_id_to_text):
                chunk_id = list(chunk_id_to_text.keys())[idx]
                retrieved_chunks.append(chunk_id_to_text[chunk_id])
        if not retrieved_chunks:
            logger.warning("No relevant chunks retrieved")
            return ["No relevant information found."]
        logger.info(f"Retrieved {len(retrieved_chunks)} chunks for query: {query[:50]}...")
        return retrieved_chunks
    except Exception as e:
        logger.error(f"Error in retrieve_relevant_chunks: {str(e)}")
        return ["Error retrieving relevant information."]

def call_llm_with_prompt(prompt_text: str) -> str:
    """Call LLM with retry logic"""
    max_retries = 3
    for attempt in range(max_retries):
        try:
            response = llm.generate_content(prompt_text)
            return response.text
        except Exception as e:
            logger.warning(f"LLM call attempt {attempt+1} failed: {str(e)}")
            if attempt == max_retries - 1:
                logger.error(f"LLM call failed after {max_retries} attempts: {str(e)}")
                st.error(f"Error calling LLM: {str(e)}")
                return ""
            time.sleep(2 ** attempt)

def parse_json_response(response_text: str, agent_type: str) -> Dict:
    """Parse and validate JSON response with jsonschema"""
    try:
        cleaned_text = response_text.strip().replace('```json', '').replace('```', '')
        data = json.loads(cleaned_text)
        schema = AGENT_SCHEMAS.get(agent_type)
        if schema:
            jsonschema.validate(instance=data, schema=schema)
        return data
    except (json.JSONDecodeError, jsonschema.ValidationError) as e:
        logger.error(f"JSON parsing/validation error for {agent_type}: {str(e)}")
        st.error(f"Error parsing JSON response: {str(e)}")
        # Fallback to default structure
        default_outputs = {
            "candidate_behavior": {
                "personality_traits": [],
                "soft_skills": [],
                "communication_patterns": [],
                "work_style": [],
                "core_values": [],
                "emotional_intelligence": "unknown",
                "summary": "Analysis failed due to invalid response."
            },
            "company_culture": {
                "core_values": [],
                "work_environment": [],
                "behavioral_expectations": [],
                "communication_norms": [],
                "preferred_work_mode": "unknown",
                "cultural_red_flags": [],
                "summary": "Analysis failed due to invalid response."
            },
            "compatibility_mapping": {
                "compatibility_scores": {},
                "alignment_strengths": [],
                "potential_mismatches": [],
                "adaptation_requirements": [],
                "summary": "Mapping failed due to invalid response."
            },
            "culture_fit_scoring": {
                "overall_fit_score": 0,
                "detailed_scores": {},
                "strengths": [],
                "challenges": [],
                "recommendations": [],
                "interview_questions": [],
                "onboarding_plan": [],
                "summary": "Scoring failed due to invalid response."
            }
        }
        return default_outputs.get(agent_type, {"error": f"JSON parsing failed: {str(e)}"})

def analyze_candidate_behavior(candidate_data: Dict[str, str]) -> Dict:
    """Agent 1: Candidate Behavior Analyzer with RAG"""
    global index, chunk_id_to_text
    index = faiss.IndexFlatL2(dimension)
    chunk_id_to_text.clear()
    
    for key, value in candidate_data.items():
        if value:
            create_chunks_and_embeddings(clean_text(value), key)
    
    if not chunk_id_to_text:
        logger.error("No valid candidate data provided")
        return {"error": "No valid candidate data provided"}
    
    query = "Extract personality traits, soft skills, communication patterns, work style, and core values from candidate data."
    retrieved_chunks = retrieve_relevant_chunks(query, k=10)
    
    prompt_text = f"""
    You are an expert behavioral analyst. Analyze the candidate's information to extract personality cues, communication patterns, and behavioral traits based on the retrieved context. Focus on evidence-based reasoning and subtle indicators like tone and word choice.

    Retrieved Context:
    {"\n".join(retrieved_chunks)}

    Provide a structured JSON output with at least 4-6 items per list:
    {{
        "personality_traits": [
            {{"trait": "openness", "evidence": "example from text", "confidence": 0.9}}
        ],
        "soft_skills": [
            {{"skill": "empathy", "evidence": "example", "level": "high", "context": "teamwork"}}
        ],
        "communication_patterns": [
            {{"pattern": "assertive", "evidence": "example", "frequency": "often"}}
        ],
        "work_style": [
            {{"style": "collaborative", "evidence": "example", "preference": "strong"}}
        ],
        "core_values": [
            {{"value": "innovation", "evidence": "example", "importance": "high"}}
        ],
        "emotional_intelligence": "high/medium/low",
        "summary": "Detailed behavioral profile summary"
    }}
    """
    response_text = call_llm_with_prompt(prompt_text)
    if response_text:
        logger.info("Candidate behavior analysis completed")
        return parse_json_response(response_text, "candidate_behavior")
    else:
        logger.error("Candidate behavior analysis failed")
        return {"error": "Analysis failed"}

def extract_company_culture(company_data: Dict[str, str]) -> Dict:
    """Agent 2: Company Culture Retriever with RAG"""
    global index, chunk_id_to_text
    index = faiss.IndexFlatL2(dimension)
    chunk_id_to_text.clear()
    
    for key, value in company_data.items():
        if value:
            create_chunks_and_embeddings(clean_text(value), key)
    
    if not chunk_id_to_text:
        logger.error("No valid company data provided")
        return {"error": "No valid company data provided"}
    
    query = "Extract core values, work environment, behavioral expectations, and communication norms from company data."
    retrieved_chunks = retrieve_relevant_chunks(query, k=10)
    
    prompt_text = f"""
    You are an expert in organizational culture analysis using RAG. Extract cultural values, behavioral expectations, and environmental traits from the provided company information.

    Retrieved Context:
    {"\n".join(retrieved_chunks)}

    Provide a structured JSON output with at least 4-6 items per list:
    {{
        "core_values": [
            {{"value": "transparency", "evidence": "example from text", "importance": "high", "source": "handbook"}}
        ],
        "work_environment": [
            {{"aspect": "agility", "evidence": "example", "expectation": "high adaptability", "source": "job_desc"}}
        ],
        "behavioral_expectations": [
            {{"behavior": "proactivity", "evidence": "example", "required_level": "high", "source": "handbook"}}
        ],
        "communication_norms": [
            {{"norm": "open feedback", "evidence": "example", "enforcement": "strict"}}
        ],
        "preferred_work_mode": "remote/hybrid/office",
        "cultural_red_flags": [
            {{"issue": "high turnover", "evidence": "example", "severity": "medium"}}
        ],
        "summary": "Comprehensive culture summary"
    }}
    """
    response_text = call_llm_with_prompt(prompt_text)
    if response_text:
        logger.info("Company culture extraction completed")
        return parse_json_response(response_text, "company_culture")
    else:
        logger.error("Company culture extraction failed")
        return {"error": "Analysis failed"}

def map_compatibility(candidate_analysis: Dict, company_analysis: Dict) -> Dict:
    """Agent 3: Compatibility Mapping"""
    if 'error' in candidate_analysis or 'error' in company_analysis:
        return {"error": "Invalid input data"}
    
    prompt_text = f"""
    Compare the candidate's behavioral traits with the company's cultural expectations.

    Candidate Profile:
    {json.dumps(candidate_analysis, indent=2)}

    Company Culture:
    {json.dumps(company_analysis, indent=2)}

    Provide a structured JSON output:
    {{
        "compatibility_scores": {{
            "values_alignment": {{"score": 90, "explanation": "Shared emphasis on innovation"}},
            "work_style_match": {{"score": 85, "explanation": "Prefers collaborative work"}},
            "communication_compatibility": {{"score": 80, "explanation": "Diplomatic style fits norms"}},
            "behavioral_fit": {{"score": 88, "explanation": "Proactivity matches expectations"}},
            "emotional_intelligence_fit": {{"score": 75, "explanation": "Moderate EI in high-pressure environment"}}
        }},
        "alignment_strengths": [
            {{"area": "teamwork", "evidence": "example", "impact": "positive"}}
        ],
        "potential_mismatches": [
            {{"area": "pace", "risk": "medium", "explanation": "Prefers structured pace"}}
        ],
        "adaptation_requirements": [
            {{"area": "agility", "adjustment": "Needs to adapt", "difficulty": "moderate"}}
        ],
        "summary": "Compatibility analysis"
    }}
    """
    response_text = call_llm_with_prompt(prompt_text)
    if response_text:
        logger.info("Compatibility mapping completed")
        return parse_json_response(response_text, "compatibility_mapping")
    else:
        logger.error("Compatibility mapping failed")
        return {"error": "Mapping failed"}

def score_culture_fit(compatibility_analysis: Dict) -> Dict:
    """Agent 4: Culture Fit Scoring"""
    if 'error' in compatibility_analysis:
        return {"error": "Invalid input data"}
    
    prompt_text = f"""
    Generate a culture fit score based on compatibility analysis.

    Compatibility Analysis:
    {json.dumps(compatibility_analysis, indent=2)}

    Provide a structured JSON output:
    {{
        "overall_fit_score": 87,
        "detailed_scores": {{
            "values": {{"score": 90, "weight": 0.3, "explanation": "Strong alignment"}},
            "work_style": {{"score": 85, "weight": 0.25, "explanation": "Good fit"}},
            "communication": {{"score": 80, "weight": 0.2, "explanation": "Compatible"}},
            "behavior": {{"score": 88, "weight": 0.15, "explanation": "Meets expectations"}},
            "emotional_intelligence": {{"score": 75, "weight": 0.1, "explanation": "Room for growth"}}
        }},
        "strengths": [
            {{"area": "collaboration", "explanation": "Team-oriented approach"}}
        ],
        "challenges": [
            {{"area": "pace", "risk_level": "medium", "explanation": "May struggle with changes"}}
        ],
        "recommendations": [
            {{"action": "Discuss pace", "context": "interview"}}
        ],
        "interview_questions": [
            {{"question": "How do you handle shifts?", "purpose": "Assess agility"}}
        ],
        "onboarding_plan": [
            {{"step": "Assign mentor", "timeline": "First 30 days"}}
        ],
        "summary": "Fit assessment"
    }}
    """
    response_text = call_llm_with_prompt(prompt_text)
    if response_text:
        logger.info("Culture fit scoring completed")
        return parse_json_response(response_text, "culture_fit_scoring")
    else:
        logger.error("Culture fit scoring failed")
        return {"error": "Scoring failed"}

def generate_bar_chart(scores_df):
    """Generate bar chart"""
    plt.figure(figsize=(8, 4))
    plt.bar(scores_df['Dimension'], scores_df['Score'], color='skyblue')
    plt.xlabel('Dimension')
    plt.ylabel('Score')
    plt.title('Score Breakdown')
    plt.xticks(rotation=45)
    plt.tight_layout()
    buffer = BytesIO()
    plt.savefig(buffer, format='png')
    buffer.seek(0)
    plt.close()
    return buffer

def generate_docx_report(candidate_analysis, company_analysis, compatibility_analysis, fit_analysis):
    """Generate DOCX report"""
    doc = Document()
    doc.add_heading('Culture Fit Analysis Report', 0)
    doc.add_paragraph(f'Date: June 19, 2025')

    if 'overall_fit_score' in fit_analysis:
        score = fit_analysis['overall_fit_score']
        doc.add_heading('Overall Culture Fit Score', level=1)
        doc.add_paragraph(f'Score: {score}/100')
        if score >= 90:
            doc.add_paragraph('Excellent Cultural Fit', style='Intense Quote')
        elif score >= 80:
            doc.add_paragraph('Good Cultural Fit', style='Intense Quote')
        elif score >= 70:
            doc.add_paragraph('Moderate Cultural Fit', style='Intense Quote')
        else:
            doc.add_paragraph('Poor Cultural Fit', style='Intense Quote')

    doc.add_heading('Candidate Profile', level=1)
    if 'summary' in candidate_analysis:
        doc.add_paragraph(candidate_analysis['summary'])
    if 'personality_traits' in candidate_analysis:
        doc.add_heading('Personality Traits', level=2)
        for trait in candidate_analysis['personality_traits'][:5]:
            doc.add_paragraph(f"• {trait.get('trait', 'N/A')}: Confidence {trait.get('confidence', 'N/A')}", style='List Bullet')

    doc.add_heading('Company Culture', level=1)
    if 'summary' in company_analysis:
        doc.add_paragraph(company_analysis['summary'])
    if 'core_values' in company_analysis:
        doc.add_heading('Core Values', level=2)
        for value in company_analysis['core_values'][:5]:
            doc.add_paragraph(f"• {value.get('value', 'N/A')}: {value.get('importance', 'N/A')} ({value.get('source', 'N/A')})", style='List Bullet')

    if 'detailed_scores' in fit_analysis:
        doc.add_heading('Score Breakdown', level=1)
        scores_df = pd.DataFrame(
            [(k, v['score'], v['weight']) for k, v in fit_analysis['detailed_scores'].items()],
            columns=['Dimension', 'Score', 'Weight']
        )
        chart_buffer = generate_bar_chart(scores_df)
        doc.add_picture(chart_buffer, width=Inches(5.0))

    doc.add_heading('Strengths', level=1)
    if 'strengths' in fit_analysis:
        for strength in fit_analysis['strengths']:
            p = doc.add_paragraph()
            p.add_run(f"{strength.get('area', 'N/A')}: ").bold = True
            p.add_run(strength.get('explanation', 'N/A'))

    doc.add_heading('Challenges', level=1)
    if 'challenges' in fit_analysis:
        for challenge in fit_analysis['challenges']:
            p = doc.add_paragraph()
            p.add_run(f"{challenge.get('area', 'N/A')}: ").bold = True
            p.add_run(challenge.get('explanation', 'N/A'))

    doc.add_heading('Recommendations', level=1)
    if 'recommendations' in fit_analysis:
        for i, rec in enumerate(fit_analysis['recommendations'], 1):
            doc.add_paragraph(f"{i}. {rec['action']} ({rec['context']})")

    doc.add_heading('Interview Questions', level=1)
    if 'interview_questions' in fit_analysis:
        for i, q in enumerate(fit_analysis['interview_questions'], 1):
            doc.add_paragraph(f"{i}. {q['question']} ({q['purpose']})")

    doc.add_heading('Onboarding Plan', level=1)
    if 'onboarding_plan' in fit_analysis:
        for i, step in enumerate(fit_analysis['onboarding_plan'], 1):
            doc.add_paragraph(f"{i}. {step['step']} ({step['timeline']})")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def display_results(candidate_analysis: Dict, company_analysis: Dict, compatibility_analysis: Dict, fit_analysis: Dict):
    """Display results"""
    st.header("📊 Culture Fit Analysis Results")
    
    docx_buffer = generate_docx_report(candidate_analysis, company_analysis, compatibility_analysis, fit_analysis)
    st.download_button(
        label="📋 Download DOCX Report",
        data=docx_buffer,
        file_name="culture_fit_report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    
    if 'overall_fit_score' in fit_analysis:
        score = fit_analysis['overall_fit_score']
        st.metric("Overall Culture Fit Score", f"{score}/100")
        if score >= 90:
            st.success("🟢 Exceptional Fit")
        elif score >= 80:
            st.info("🔵 Strong Fit")
        elif score >= 70:
            st.warning("🟡 Moderate Fit")
        else:
            st.error("🔴 Limited Fit")
    
    col1, col2 = st.columns(2)
    with col1:
        st.subheader("👤 Candidate Profile")
        if 'summary' in candidate_analysis:
            st.write(candidate_analysis['summary'])
        if 'personality_traits' in candidate_analysis:
            st.write("**Key Traits:**")
            for trait in candidate_analysis['personality_traits'][:5]:
                st.write(f"• {trait.get('trait', 'N/A')}: Confidence {trait.get('confidence', 'N/A')}")
    
    with col2:
        st.subheader("🏢 Company Culture")
        if 'summary' in company_analysis:
            st.write(company_analysis['summary'])
        if 'core_values' in company_analysis:
            st.write("**Core Values:**")
            for value in company_analysis['core_values'][:5]:
                st.write(f"• {value.get('value', 'N/A')}: {value.get('importance', 'N/A')} ({value.get('source', 'N/A')})")
    
    if 'detailed_scores' in fit_analysis:
        st.subheader("📈 Score Breakdown")
        scores_df = pd.DataFrame(
            [(k, v['score']) for k, v in fit_analysis['detailed_scores'].items()],
            columns=['Dimension', 'Score']
        )
        st.bar_chart(scores_df.set_index('Dimension')['Score'])
    
    col3, col4 = st.columns(2)
    with col3:
        if 'strengths' in fit_analysis:
            st.subheader("✅ Strengths")
            for strength in fit_analysis['strengths']:
                st.success(f"**{strength.get('area', 'N/A')}**: {strength.get('explanation', 'N/A')}")
    
    with col4:
        if 'challenges' in fit_analysis:
            st.subheader("⚠️ Challenges")
            for challenge in fit_analysis['challenges']:
                risk_level = challenge.get('risk_level', 'medium')
                if risk_level == 'high':
                    st.error(f"**{challenge.get('area', 'N/A')}**: {challenge.get('explanation', 'N/A')}")
                elif risk_level == 'medium':
                    st.warning(f"**{challenge.get('area', 'N/A')}**: {challenge.get('explanation', 'N/A')}")
                else:
                    st.info(f"**{challenge.get('area', 'N/A')}**: {challenge.get('explanation', 'N/A')}")
    
    col5, col6 = st.columns(2)
    with col5:
        if 'recommendations' in fit_analysis:
            st.subheader("💡 Recommendations")
            for i, rec in enumerate(fit_analysis['recommendations'], 1):
                st.write(f"{i}. {rec['action']} ({rec['context']})")
    
    with col6:
        if 'interview_questions' in fit_analysis:
            st.subheader("🎯 Interview Questions")
            for i, q in enumerate(fit_analysis['interview_questions'], 1):
                st.write(f"{i}. {q['question']} ({q['purpose']})")
    
    if 'onboarding_plan' in fit_analysis:
        st.subheader("🚀 Onboarding Plan")
        for i, step in enumerate(fit_analysis['onboarding_plan'], 1):
            st.write(f"{i}. {step['step']} ({step['timeline']})")

# Streamlit UI
st.set_page_config(page_title="Culture Fit Analyzer", layout="wide")
st.title("🧠 AI-Powered Culture Fit Analyzer")
st.markdown("Evaluate candidate alignment with company culture using advanced AI insights.")

with st.sidebar:
    st.header("📋 Instructions")
    st.markdown("""
    **Quick Analysis:**
    1. Enter candidate profile URLs (e.g., LinkedIn, GitHub)
    2. Enter company URLs (e.g., About, Glassdoor)
    3. Click 'Quick Analysis'
                
    **Manual Input:**
    1. Upload candidate files or enter text
    2. Upload company files or enter text
    3. Click 'Analyze Culture Fit'
    
    **Results:**
    - View results in the 'Results' tab
    - Download report as DOCX
    """)

tab1, tab2, tab3 = st.tabs(["⚡ Quick Analysis", "📄 Manual Input", "📊 Results"])

with tab1:
    st.header("⚡ Quick Analysis - URL Based")
    col1, col2 = st.columns(2)
    
    with col1:
        st.subheader("👤 Candidate URLs")
        candidate_linkedin = st.text_input("LinkedIn Profile URL", placeholder="https://www.linkedin.com/in/username", value="https://www.linkedin.com/in/afrin-s-snsihub-412139370/")
        candidate_github = st.text_input("GitHub Profile URL", placeholder="https://github.com/username")
        candidate_portfolio = st.text_input("Portfolio/Personal Website", placeholder="https://portfolio.com",value="https://brendaneich.com/")
        candidate_other = st.text_input("Other Profile URL", placeholder="Any other relevant URL", value="https://news.ycombinator.com/user?id=torvalds")
    
    with col2:
        st.subheader("🏢 Company URLs")
        company_about = st.text_input("Company About Page", placeholder="https://company.com/about", value="https://about.gitlab.com/jobs/")
        company_careers = st.text_input("Careers/Jobs Page", placeholder="https://company.com/careers")
        company_glassdoor = st.text_input("Glassdoor URL", placeholder="https://glassdoor.com/company/...")
        company_other = st.text_input("Other Company URL", placeholder="Any other relevant URL", value="https://handbook.gitlab.com/handbook/values/")
    
    if st.button("🚀 Quick Analysis", type="primary"):
        candidate_urls = [url for url in [candidate_linkedin, candidate_github, candidate_portfolio, candidate_other] if url.strip()]
        company_urls = [url for url in [company_about, company_careers, company_glassdoor, company_other] if url.strip()]
        
        if not candidate_urls or not company_urls:
            st.error("Provide at least one candidate URL and one company URL.")
        else:
            with st.spinner("Scraping and analyzing content..."):
                candidate_content = {}
                for i, url in enumerate(candidate_urls):
                    with st.status(f"Scraping candidate URL {i+1}/{len(candidate_urls)}...") as status:
                        content = scrape_website_content(url)
                        if content:
                            st.write(f"**Preview:** {content[:200]}{'...' if len(content) > 200 else ''}")
                            if 'linkedin' in url.lower():
                                candidate_content['linkedin'] = content
                            elif 'github' in url.lower():
                                candidate_content['github'] = content
                            else:
                                candidate_content['bio'] = candidate_content.get('bio', '') + content
                        else:
                            st.write("No content scraped from this URL.")
                        status.update(label=f"Completed scraping candidate URL {i+1}/{len(candidate_urls)}")
                
                company_content = {}
                for i, url in enumerate(company_urls):
                    with st.status(f"Scraping company URL {i+1}/{len(company_urls)}...") as status:
                        content = scrape_website_content(url)
                        if content:
                            st.write(f"**Preview:** {content[:200]}{'...' if len(content) > 200 else ''}")
                            if 'about' in url.lower() or 'company' in url.lower():
                                company_content['about_us'] = content
                            elif 'career' in url.lower() or 'job' in url.lower():
                                company_content['job_desc'] = content
                            elif 'glassdoor' in url.lower():
                                company_content['reviews'] = content
                            else:
                                company_content['about_us'] = company_content.get('about_us', '') + content
                        else:
                            st.write("No content scraped from this URL.")
                        status.update(label=f"Completed scraping company URL {i+1}/{len(company_urls)}")
                
                manual_candidate_data = st.session_state.get('manual_candidate_data', {})
                manual_company_data = st.session_state.get('manual_company_data', {})
                
                if not candidate_content and not manual_candidate_data:
                    st.error("No valid candidate data. Use 'Manual Input' tab.")
                elif not company_content and not manual_company_data:
                    st.error("No valid company data. Use 'Manual Input' tab.")
                else:
                    with st.status("Analyzing candidate behavior...") as status:
                        combined_candidate_data = {**candidate_content, **manual_candidate_data}
                        combined_company_data = {**company_content, **manual_company_data}
                        
                        candidate_analysis = analyze_candidate_behavior(combined_candidate_data)
                        status.update(label="Extracting company culture...")
                        
                        if 'error' not in candidate_analysis:
                            company_analysis = extract_company_culture(combined_company_data)
                            status.update(label="Mapping compatibility...")
                            
                            if 'error' not in company_analysis:
                                compatibility_analysis = map_compatibility(candidate_analysis, company_analysis)
                                status.update(label="Scoring culture fit...")
                                
                                if 'error' not in compatibility_analysis:
                                    fit_analysis = score_culture_fit(compatibility_analysis)
                                    status.update(label="Analysis complete!", state="complete")
                                    
                                    st.session_state['candidate_analysis'] = candidate_analysis
                                    st.session_state['company_analysis'] = company_analysis
                                    st.session_state['compatibility_analysis'] = compatibility_analysis
                                    st.session_state['fit_analysis'] = fit_analysis
                                    
                                    st.success("Analysis complete! Check the Results tab.")
                                else:
                                    st.error("Compatibility mapping failed.")
                            else:
                                st.error("Company analysis failed.")
                        else:
                            st.error("Candidate profiling failed.")

with tab2:
    col1, col2 = st.columns(2)
    
    with col1:
        st.header("👤 Candidate Information")
        resume_file = st.file_uploader("Upload Resume/CV (PDF)", type="pdf")
        linkedin_bio = st.text_area("LinkedIn Profile / Professional Bio", height=120)
        github_content = st.text_area("GitHub README / Portfolio Description", height=100)
        personal_bio = st.text_area("Personal Statement / Cover Letter", height=100)
    
    with col2:
        st.header("🏢 Company Information")
        hr_file = st.file_uploader("Upload HR Policy / Handbook (PDF)", type="pdf")
        about_company = st.text_area("About Us / Company Description", height=120)
        employee_reviews = st.text_area("Employee Reviews / Glassdoor Summary", height=100)
        job_description = st.text_area("Job Description", height=100)
    
    if st.button("🔍 Analyze Culture Fit", type="primary"):
        with st.spinner("Running AI analysis..."):
            candidate_data = {
                'resume': extract_text_from_pdf(resume_file) if resume_file else "",
                'linkedin': linkedin_bio or "",
                'github': github_content or "",
                'bio': personal_bio or ""
            }
            
            company_data = {
                'hr_policy': extract_text_from_pdf(hr_file) if hr_file else "",
                'about_us': about_company or "",
                'reviews': employee_reviews or "",
                'job_desc': job_description or ""
            }
            
            if not any(candidate_data.values()) or not any(company_data.values()):
                st.error("Provide at least some candidate and company information.")
            else:
                with st.status("Analyzing candidate behavior...") as status:
                    st.session_state['manual_candidate_data'] = candidate_data
                    st.session_state['manual_company_data'] = company_data
                    
                    candidate_analysis = analyze_candidate_behavior(candidate_data)
                    status.update(label="Extracting company culture...")
                    
                    if 'error' not in candidate_analysis:
                        company_analysis = extract_company_culture(company_data)
                        status.update(label="Mapping compatibility...")
                        
                        if 'error' not in company_analysis:
                            compatibility_analysis = map_compatibility(candidate_analysis, company_analysis)
                            status.update(label="Scoring culture fit...")
                            
                            if 'error' not in compatibility_analysis:
                                fit_analysis = score_culture_fit(compatibility_analysis)
                                status.update(label="Analysis complete!", state="complete")
                                
                                st.session_state['candidate_analysis'] = candidate_analysis
                                st.session_state['company_analysis'] = company_analysis
                                st.session_state['compatibility_analysis'] = compatibility_analysis
                                st.session_state['fit_analysis'] = fit_analysis
                                
                                st.success("Analysis complete! Check the Results tab.")
                            else:
                                st.error("Compatibility mapping failed.")
                        else:
                            st.error("Company analysis failed.")
                    else:
                        st.error("Candidate analysis failed.")

with tab3:
    if all(key in st.session_state for key in ['candidate_analysis', 'company_analysis', 'compatibility_analysis', 'fit_analysis']):
        display_results(
            st.session_state['candidate_analysis'],
            st.session_state['company_analysis'],
            st.session_state['compatibility_analysis'],
            st.session_state['fit_analysis']
        )
    else:
        st.info("Run 'Quick Analysis' or 'Manual Analysis' first to see results.")

st.markdown("---")
st.markdown("*This tool uses AI to provide culture fit insights. Results should complement human judgment.*")